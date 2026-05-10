import re
import os
from backend.schemas import Textbook, Chapter, TextbookStatus
from backend import config


def parse_file(textbook: Textbook) -> Textbook:
    file_path = os.path.join(config.UPLOAD_DIR, f"{textbook.textbook_id}.{textbook.file_type}")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    if textbook.file_type == "pdf":
        return _parse_pdf(file_path, textbook)
    elif textbook.file_type in ("md", "markdown"):
        return _parse_markdown(file_path, textbook)
    elif textbook.file_type == "txt":
        return _parse_txt(file_path, textbook)
    elif textbook.file_type == "docx":
        return _parse_docx(file_path, textbook)
    else:
        raise ValueError(f"Unsupported file type: {textbook.file_type}")


def _parse_pdf(file_path: str, textbook: Textbook) -> Textbook:
    import pymupdf

    try:
        doc = pymupdf.open(file_path)
    except Exception as e:
        raise ValueError(f"无法解析PDF文件，可能已损坏: {e}")
    total_pages = doc.page_count

    pages_text = []
    for i in range(total_pages):
        page = doc[i]
        text = page.get_text("text", sort=True)
        pages_text.append((i + 1, text))

    doc.close()

    full_text = "\n".join([t for _, t in pages_text])
    total_chars = len(full_text)

    chapters = _detect_chapters(pages_text, textbook.textbook_id)

    if not chapters:
        chapters = _create_virtual_chapters(pages_text, textbook.textbook_id)

    textbook.total_pages = total_pages
    textbook.total_chars = sum(c.char_count for c in chapters)
    textbook.chapters = chapters

    return textbook


CHAPTER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百零0-9]+章\s*[^\n]{0,60}", re.MULTILINE),
    re.compile(r"^第[一二三四五六七八九十百零0-9]+节\s*[^\n]{0,60}", re.MULTILINE),
    re.compile(r"^Chapter\s+\d+[:\s]?.*", re.MULTILINE | re.IGNORECASE),
    re.compile(r"^\d+\.\d*\s+[^\n]{2,60}$", re.MULTILINE),
]


def _detect_chapters(pages_text: list[tuple[int, str]], textbook_id: str) -> list[Chapter]:
    all_text = ""
    page_map = []
    for page_num, text in pages_text:
        offset = len(all_text)
        all_text += text + "\n"
        page_map.append((offset, page_num))

    def offset_to_page(offset: int) -> int:
        result = 1
        for start, pn in page_map:
            if offset >= start:
                result = pn
            else:
                break
        return result

    matches = []
    for pattern in CHAPTER_PATTERNS:
        for m in pattern.finditer(all_text):
            title = m.group().strip()
            if len(title) >= 3:
                matches.append((m.start(), title))

    matches.sort(key=lambda x: x[0])

    if len(matches) < 2:
        return []

    chapters = []
    for i, (start, title) in enumerate(matches):
        end = matches[i + 1][0] if i + 1 < len(matches) else len(all_text)
        content = all_text[start:end].strip()
        page_start = offset_to_page(start)
        page_end = offset_to_page(end - 1) if end > start else page_start

        chapters.append(Chapter(
            chapter_id=f"{textbook_id}_ch_{i:02d}",
            textbook_id=textbook_id,
            title=title[:80],
            page_start=page_start,
            page_end=page_end,
            content=content,
            char_count=len(content),
        ))

    return chapters


def _create_virtual_chapters(pages_text: list[tuple[int, str]], textbook_id: str) -> list[Chapter]:
    chapters = []
    pages_per_chapter = 10

    for i in range(0, len(pages_text), pages_per_chapter):
        chunk = pages_text[i:i + pages_per_chapter]
        content = "\n".join([t for _, t in chunk])
        page_start = chunk[0][0]
        page_end = chunk[-1][0]

        chapters.append(Chapter(
            chapter_id=f"{textbook_id}_ch_{len(chapters):02d}",
            textbook_id=textbook_id,
            title=f"第{len(chapters) + 1}部分 (第{page_start}-{page_end}页)",
            page_start=page_start,
            page_end=page_end,
            content=content,
            char_count=len(content),
        ))

    return chapters


def _parse_markdown(file_path: str, textbook: Textbook) -> Textbook:
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    heading_pattern = re.compile(r"^(#{1,3})\s+(.+)$", re.MULTILINE)
    matches = list(heading_pattern.finditer(content))

    chapters = []
    if matches:
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            ch_content = content[start:end].strip()
            title = m.group(2).strip()

            chapters.append(Chapter(
                chapter_id=f"{textbook.textbook_id}_ch_{i:02d}",
                textbook_id=textbook.textbook_id,
                title=title[:80],
                page_start=i + 1,
                page_end=i + 1,
                content=ch_content,
                char_count=len(ch_content),
            ))
    else:
        chapters.append(Chapter(
            chapter_id=f"{textbook.textbook_id}_ch_00",
            textbook_id=textbook.textbook_id,
            title="全文",
            page_start=1,
            page_end=1,
            content=content,
            char_count=len(content),
        ))

    textbook.total_pages = len(chapters)
    textbook.total_chars = sum(c.char_count for c in chapters)
    textbook.chapters = chapters
    return textbook


def _parse_txt(file_path: str, textbook: Textbook) -> Textbook:
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    chapters = []
    lines = content.split("\n")
    current_title = ""
    current_content = []
    chapter_idx = 0

    for line in lines:
        stripped = line.strip()
        if re.match(r"^第[一二三四五六七八九十百零0-9]+章", stripped) or \
           re.match(r"^第[一二三四五六七八九十百零0-9]+节", stripped):
            if current_content:
                ch_text = "\n".join(current_content)
                chapters.append(Chapter(
                    chapter_id=f"{textbook.textbook_id}_ch_{chapter_idx:02d}",
                    textbook_id=textbook.textbook_id,
                    title=current_title[:80] or f"第{chapter_idx + 1}部分",
                    page_start=chapter_idx + 1,
                    page_end=chapter_idx + 1,
                    content=ch_text,
                    char_count=len(ch_text),
                ))
                chapter_idx += 1
            current_title = stripped
            current_content = [line]
        else:
            current_content.append(line)

    if current_content:
        ch_text = "\n".join(current_content)
        chapters.append(Chapter(
            chapter_id=f"{textbook.textbook_id}_ch_{chapter_idx:02d}",
            textbook_id=textbook.textbook_id,
            title=current_title[:80] or f"第{chapter_idx + 1}部分",
            page_start=chapter_idx + 1,
            page_end=chapter_idx + 1,
            content=ch_text,
            char_count=len(ch_text),
        ))

    if not chapters:
        chapters.append(Chapter(
            chapter_id=f"{textbook.textbook_id}_ch_00",
            textbook_id=textbook.textbook_id,
            title="全文",
            page_start=1,
            page_end=1,
            content=content,
            char_count=len(content),
        ))

    textbook.total_pages = len(chapters)
    textbook.total_chars = sum(c.char_count for c in chapters)
    textbook.chapters = chapters
    return textbook


def _parse_docx(file_path: str, textbook: Textbook) -> Textbook:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)

    chapters = []
    current_title = ""
    current_content = []
    chapter_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        is_heading = (
            para.style.name.startswith("Heading") or
            re.match(r"^第[一二三四五六七八九十百零0-9]+章", text) or
            re.match(r"^第[一二三四五六七八九十百零0-9]+节", text)
        )

        if is_heading:
            if current_content:
                ch_text = "\n".join(current_content)
                chapters.append(Chapter(
                    chapter_id=f"{textbook.textbook_id}_ch_{chapter_idx:02d}",
                    textbook_id=textbook.textbook_id,
                    title=current_title[:80] or f"第{chapter_idx + 1}部分",
                    page_start=chapter_idx + 1,
                    page_end=chapter_idx + 1,
                    content=ch_text,
                    char_count=len(ch_text),
                ))
                chapter_idx += 1
            current_title = text
            current_content = [text]
        else:
            current_content.append(text)

    if current_content:
        ch_text = "\n".join(current_content)
        chapters.append(Chapter(
            chapter_id=f"{textbook.textbook_id}_ch_{chapter_idx:02d}",
            textbook_id=textbook.textbook_id,
            title=current_title[:80] or f"第{chapter_idx + 1}部分",
            page_start=chapter_idx + 1,
            page_end=chapter_idx + 1,
            content=ch_text,
            char_count=len(ch_text),
        ))

    if not chapters:
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        chapters.append(Chapter(
            chapter_id=f"{textbook.textbook_id}_ch_00",
            textbook_id=textbook.textbook_id,
            title="全文",
            page_start=1,
            page_end=1,
            content=full_text,
            char_count=len(full_text),
        ))

    textbook.total_pages = len(chapters)
    textbook.total_chars = sum(c.char_count for c in chapters)
    textbook.chapters = chapters
    return textbook
